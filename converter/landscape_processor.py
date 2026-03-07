"""Post-processing for mixed portrait/landscape sections in DOCX files.

Inserts section breaks and switches page orientation for specified sections,
while preserving header/footer references across all sections.
"""

import logging
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)

EMU_PER_TWIP = 635

class LandscapeSection:
    """Defines a section of the document that should be in landscape orientation.

    Attributes:
        start_marker: Text pattern that marks the beginning of the landscape section.
        end_marker: Text pattern that marks the end of the landscape section.
        narrow_margins: Whether to use narrow margins (12.7mm) in landscape. Default True.
    """

    def __init__(
        self,
        start_marker: str,
        end_marker: str,
        narrow_margins: bool = True,
    ):
        self.start_marker = start_marker
        self.end_marker = end_marker
        self.narrow_margins = narrow_margins


def _find_paragraph_index(doc: Document, marker: str) -> Optional[int]:
    """Find the index of a paragraph containing the marker text."""
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            return i
    return None


def _copy_header_footer_refs(
    source_sectPr, target_sectPr
) -> Tuple[int, int]:
    """Copy header and footer references from source to target section.
    Returns (header_count, footer_count)."""
    hdr_refs = source_sectPr.findall(qn('w:headerReference'))
    ftr_refs = source_sectPr.findall(qn('w:footerReference'))
    for h in hdr_refs:
        target_sectPr.append(etree.fromstring(etree.tostring(h)))
    for f in ftr_refs:
        target_sectPr.append(etree.fromstring(etree.tostring(f)))
    return len(hdr_refs), len(ftr_refs)


def _insert_section_break(
    paragraph_element,
    page_w: int,
    page_h: int,
    margins: dict,
    orientation: str,
    header_footer_refs,
) -> None:
    """Insert a section break before the given paragraph element."""
    pPr = paragraph_element.get_or_add_pPr()
    # Remove existing sectPr
    for existing in pPr.findall(qn('w:sectPr')):
        pPr.remove(existing)

    sectPr = etree.SubElement(pPr, qn('w:sectPr'))
    etree.SubElement(sectPr, qn('w:type')).set(qn('w:val'), 'nextPage')

    pgSz = etree.SubElement(sectPr, qn('w:pgSz'))
    if orientation == 'landscape':
        pgSz.set(qn('w:w'), str(page_h))  # swap
        pgSz.set(qn('w:h'), str(page_w))
        pgSz.set(qn('w:orient'), 'landscape')
    else:
        pgSz.set(qn('w:w'), str(page_w))
        pgSz.set(qn('w:h'), str(page_h))

    pgMar = etree.SubElement(sectPr, qn('w:pgMar'))
    pgMar.set(qn('w:top'), str(margins['top']))
    pgMar.set(qn('w:bottom'), str(margins['bottom']))
    pgMar.set(qn('w:left'), str(margins['left']))
    pgMar.set(qn('w:right'), str(margins['right']))

    # Copy header/footer references
    hdr_refs, ftr_refs = header_footer_refs
    for h in hdr_refs:
        sectPr.append(etree.fromstring(etree.tostring(h)))
    for f in ftr_refs:
        sectPr.append(etree.fromstring(etree.tostring(f)))


def apply_landscape_sections(
    docx_path: Path,
    sections: List[LandscapeSection],
) -> bool:
    """Apply landscape orientation to specified sections of a DOCX file.

    Finds paragraphs matching start/end markers, inserts section breaks,
    and switches orientation to landscape. Preserves header/footer refs.

    Args:
        docx_path: Path to the DOCX file to modify (in-place).
        sections: List of LandscapeSection definitions.

    Returns:
        True if any landscape sections were applied, False otherwise.
    """
    if not sections:
        return False

    doc = Document(str(docx_path))
    applied = False

    # Get page dimensions from first section (in twips)
    sec0 = doc.sections[0]
    page_w = int(sec0.page_width / EMU_PER_TWIP)
    page_h = int(sec0.page_height / EMU_PER_TWIP)
    portrait_margins = {
        'top': int(sec0.top_margin / EMU_PER_TWIP),
        'bottom': int(sec0.bottom_margin / EMU_PER_TWIP),
        'left': int(sec0.left_margin / EMU_PER_TWIP),
        'right': int(sec0.right_margin / EMU_PER_TWIP),
    }
    narrow = int(Mm(12.7) / EMU_PER_TWIP)
    narrow_margins = {
        'top': narrow, 'bottom': narrow,
        'left': narrow, 'right': narrow,
    }

    # Collect header/footer references from first section
    body_sectPr = sec0._sectPr
    hdr_refs = body_sectPr.findall(qn('w:headerReference'))
    ftr_refs = body_sectPr.findall(qn('w:footerReference'))
    hdr_ftr = (hdr_refs, ftr_refs)

    for ls in sections:
        start_idx = _find_paragraph_index(doc, ls.start_marker)
        end_idx = _find_paragraph_index(doc, ls.end_marker)

        if start_idx is None:
            logger.warning(
                f"Landscape start marker not found: '{ls.start_marker}'"
            )
            continue
        if end_idx is None:
            logger.warning(
                f"Landscape end marker not found: '{ls.end_marker}'"
            )
            continue
        if start_idx >= end_idx:
            logger.warning(
                f"Start marker ({start_idx}) must be before "
                f"end marker ({end_idx})"
            )
            continue

        logger.info(
            f"Landscape section: paragraphs {start_idx}-{end_idx}"
        )

        # Section break BEFORE landscape (ends portrait)
        prev_elem = doc.paragraphs[start_idx - 1]._element
        _insert_section_break(
            prev_elem, page_w, page_h,
            portrait_margins, 'portrait', hdr_ftr,
        )

        # Section break AFTER landscape (ends landscape)
        # Re-find end_idx since paragraph indices haven't changed
        prev_end = doc.paragraphs[end_idx - 1]._element
        land_margins = narrow_margins if ls.narrow_margins else portrait_margins
        _insert_section_break(
            prev_end, page_w, page_h,
            land_margins, 'landscape', hdr_ftr,
        )

        applied = True
        logger.info("Landscape section breaks inserted")

    if applied:
        doc.save(str(docx_path))
        n_sections = len(doc.sections)
        logger.info(
            f"Saved with {n_sections} sections: {docx_path.name}"
        )

    return applied
